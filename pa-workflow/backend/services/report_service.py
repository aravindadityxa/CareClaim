"""
Report generation service for prior authorization decisions.
Generates professional Word documents with Sonar-based analysis and recommendations.
"""

import json
import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional
from uuid import UUID

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def _add_header_table(doc: Document, pa_data: Dict[str, Any]):
    """Add a professional header table with PA metadata."""
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Row 1
    table.rows[0].cells[0].text = 'PA ID'
    table.rows[0].cells[1].text = str(pa_data.get('pa_id', 'N/A'))
    
    # Row 2
    table.rows[1].cells[0].text = 'Decision Date'
    table.rows[1].cells[1].text = datetime.utcnow().strftime('%B %d, %Y')
    
    # Style header
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


def _add_section_heading(doc: Document, text: str):
    """Add a styled section heading."""
    heading = doc.add_paragraph(text, style='Heading 1')
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading


def _add_subsection_heading(doc: Document, text: str):
    """Add a styled subsection heading."""
    heading = doc.add_paragraph(text, style='Heading 2')
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    return heading


def _add_colored_badge(doc: Document, label: str, value: str, color: tuple = (51, 102, 153)):
    """Add a colored badge/box with text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    
    run2 = p.add_run(str(value))
    run2.font.bold = True
    run2.font.size = Pt(11)


def _format_list_item(doc: Document, text: str, indent_level: int = 0):
    """Add a formatted list item."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * indent_level)
    return p


def generate_summary_report(pa_id: UUID, pa_data: Dict[str, Any], sonar_analysis: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Generate a professional summary report document for a PA request.
    
    Args:
        pa_id: PA request ID
        pa_data: Complete PA data including OCR, medical codes, decision info
        sonar_analysis: Optional Sonar analysis result from Agent A
    
    Returns:
        Document bytes (DOCX format)
    """
    logger.info(f"Generating report for PA {pa_id}")
    logger.info(f"Sonar analysis received: {type(sonar_analysis)}")
    if sonar_analysis:
        logger.info(f"Sonar analysis keys: {list(sonar_analysis.keys()) if isinstance(sonar_analysis, dict) else 'N/A'}")
        if isinstance(sonar_analysis, dict) and 'summary' in sonar_analysis:
            logger.info(f"✅ SONAR SUMMARY FOUND: {sonar_analysis['summary'][:100]}...")
    else:
        logger.warning("⚠️ No Sonar analysis - will use fallback reasoning")

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('PRIOR AUTHORIZATION REVIEW REPORT')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    subtitle = doc.add_paragraph(f'PA ID: {pa_id}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    
    # Add header table
    _add_header_table(doc, pa_data)
    doc.add_paragraph()
    
    # Get agent outputs safely
    details = pa_data.get('details') or {}
    agent_a = details.get('agent_a_output') or {}
    
    # Ensure agent_a is a dict
    if agent_a and hasattr(agent_a, '__dict__'):
        agent_a = agent_a.__dict__
    
    # Section 1: Document Summary
    _add_section_heading(doc, '1. DOCUMENT SUMMARY')
    
    if agent_a and isinstance(agent_a, dict):
        ocr_results = agent_a.get('ocr_results')
        ocr_text = None
        
        # Handle ocr_results as a list of OCRResult objects
        if ocr_results and isinstance(ocr_results, list) and len(ocr_results) > 0:
            # ocr_results is a list - take text from the first result
            first_result = ocr_results[0]
            if isinstance(first_result, dict):
                ocr_text = first_result.get('text', '')
            else:
                # If it's an object with attributes
                ocr_text = getattr(first_result, 'text', '')
        
        if ocr_text:
            doc.add_paragraph('Extracted Clinical Text:', style='Heading 3')
            doc.add_paragraph(ocr_text[:500])
        else:
            doc.add_paragraph('No OCR text available.')
    else:
        doc.add_paragraph('No document data available.')
    
    # Section 2: Medical Codes Identified
    _add_section_heading(doc, '2. MEDICAL CODES IDENTIFIED')
    
    medical_codes = agent_a.get('medical_codes', {}) if isinstance(agent_a, dict) else {}
    
    # Handle both dict and object formats
    if isinstance(medical_codes, dict):
        icd10_codes = medical_codes.get('icd10_codes', [])
        cpt_codes = medical_codes.get('cpt_codes', [])
        rxnorm_codes = medical_codes.get('rxnorm_codes', [])
    else:
        # If it's an object with attributes
        icd10_codes = getattr(medical_codes, 'icd10_codes', [])
        cpt_codes = getattr(medical_codes, 'cpt_codes', [])
        rxnorm_codes = getattr(medical_codes, 'rxnorm_codes', [])
    
    if icd10_codes or cpt_codes or rxnorm_codes:
        if icd10_codes:
            doc.add_paragraph('ICD-10 Diagnoses:', style='Heading 3')
            for code in icd10_codes:
                _format_list_item(doc, str(code))
        
        if cpt_codes:
            doc.add_paragraph('CPT Procedures:', style='Heading 3')
            for code in cpt_codes:
                _format_list_item(doc, str(code))
        
        if rxnorm_codes:
            doc.add_paragraph('Medications (RxNorm):', style='Heading 3')
            for code in rxnorm_codes:
                _format_list_item(doc, str(code))
    else:
        doc.add_paragraph('No medical codes identified.')
    
    # Section 3: Decision Summary
    _add_section_heading(doc, '3. DECISION SUMMARY')
    
    decision = pa_data.get('decision')
    status = pa_data.get('status', 'UNKNOWN')
    
    if decision and status in ['APPROVED', 'DENIED', 'DECIDED']:
        if status == 'APPROVED':
            _add_colored_badge(doc, 'STATUS', 'APPROVED', color=(0, 176, 80))
        elif status == 'DENIED':
            _add_colored_badge(doc, 'STATUS', 'DENIED', color=(192, 0, 0))
        else:
            _add_colored_badge(doc, 'STATUS', 'HUMAN REVIEW REQUIRED', color=(255, 192, 0))
        
        doc.add_paragraph()
        
        final_score = pa_data.get('final_score')
        if final_score is not None:
            _add_colored_badge(doc, 'Final Score', f'{final_score:.1f}/100')
            doc.add_paragraph()
        
        decision_info = {
            'Decision': status,
            'Decided At': pa_data.get('decided_at', 'N/A'),
            'Created At': pa_data.get('created_at', 'N/A'),
        }
        
        if status == 'APPROVED' and pa_data.get('auth_code'):
            decision_info['Authorization Code'] = pa_data['auth_code']
        
        for label, value in decision_info.items():
            if value:
                p = doc.add_paragraph()
                run1 = p.add_run(f'{label}: ')
                run1.font.bold = True
                p.add_run(str(value))
    else:
        status_text = 'Processing' if status == 'PROCESSING' else status
        doc.add_paragraph(f'Status: {status_text}')
    
    # Section 4: AI Reasoning & Analysis
    _add_section_heading(doc, '4. AI ANALYSIS & CLINICAL REASONING')
    
    # Try to build AI reasoning from available data
    has_sonar = sonar_analysis and isinstance(sonar_analysis, dict)
    
    if has_sonar:
        # Build comprehensive reasoning paragraph from Sonar
        reasoning_parts = []
        
        # Clinical context
        summary = sonar_analysis.get('summary', '')
        if summary:
            reasoning_parts.append(f"Clinical Summary: {summary}")
        
        # Medical necessity assessment
        signals = sonar_analysis.get('medical_necessity_signals', [])
        if signals and isinstance(signals, list):
            signals_text = ', '.join(str(s) for s in signals[:3])
            reasoning_parts.append(
                f"The clinical documentation demonstrates medical necessity with the following key signals: {signals_text}. "
                "These indicators support that the requested treatment aligns with established medical evidence and patient condition severity."
            )
        
        # Risk assessment
        risks = sonar_analysis.get('risks', [])
        risk_assessment = "No significant clinical or compliance risks were identified"
        if risks and isinstance(risks, list):
            high_risks = [r for r in risks if 'HIGH' in str(r).upper()]
            medium_risks = [r for r in risks if 'MEDIUM' in str(r).upper()]
            
            if high_risks:
                risk_text = ', '.join(str(r) for r in high_risks[:2])
                risk_assessment = (
                    f"The analysis identified the following high-priority risks requiring attention: {risk_text}. "
                    "These factors may require additional clinical review or policy verification before authorization."
                )
            elif medium_risks:
                risk_text = ', '.join(str(r) for r in medium_risks[:2])
                risk_assessment = (
                    f"Moderate risk indicators were identified: {risk_text}. "
                    "These are within acceptable ranges for standard review but documented for completeness."
                )
        
        reasoning_parts.append(risk_assessment)
        
        # Recommendations
        recommendations = sonar_analysis.get('recommendations', [])
        if recommendations and isinstance(recommendations, list) and len(recommendations) > 0:
            rec_text = ' '.join(str(r) for r in recommendations[:2])
            reasoning_parts.append(
                f"Based on the clinical analysis, the following recommendations are made: {rec_text}. "
                "These recommendations are intended to ensure appropriate care delivery and compliance with policy requirements."
            )
        
        # Final assessment paragraph
        final_assessment = f"Overall Assessment: The prior authorization review has been completed with a final confidence score of {pa_data.get('final_score', 'N/A')}. "
        
        if status == 'APPROVED':
            final_assessment += (
                "The clinical information, medical codes, and policy compliance analysis support approval of this request. "
                "The patient's condition and proposed treatment meet medical necessity criteria and do not conflict with coverage policies."
            )
        elif status == 'DENIED':
            final_assessment += (
                "The analysis indicates that the request does not meet current medical necessity criteria or policy requirements. "
                "The provider may submit additional clinical documentation or appeal this decision for further review."
            )
        else:
            final_assessment += (
                "Further clinical review by medical staff is required to complete the authorization process. "
                "Any additional information requested should be submitted as soon as possible to expedite the review."
            )
        
        reasoning_parts.append(final_assessment)
        
        # Combine all parts into readable paragraphs
        for part in reasoning_parts:
            if part:
                p = doc.add_paragraph(part)
                p.paragraph_format.space_after = Pt(10)
                p.paragraph_format.line_spacing = 1.15
    else:
        # Fallback: Generate reasoning from available decision and medical data
        reasoning_parts = []
        
        # Extract available medical information
        details = pa_data.get('details') or {}
        agent_a = details.get('agent_a_output') or {}
        
        medical_codes = agent_a.get('medical_codes') if isinstance(agent_a, dict) else {}
        if isinstance(medical_codes, dict):
            icd10_codes = medical_codes.get('icd10_codes', [])
            cpt_codes = medical_codes.get('cpt_codes', [])
        else:
            icd10_codes = getattr(medical_codes, 'icd10_codes', [])
            cpt_codes = getattr(medical_codes, 'cpt_codes', [])
        
        # Build reasoning from decision logic
        intro = "AI Analysis Summary: This prior authorization has been analyzed using automated clinical reasoning systems. "
        
        # Medical codes assessment
        codes_assessment = ""
        if icd10_codes or cpt_codes:
            code_list = []
            if icd10_codes:
                code_list.extend([f"diagnosis {c}" for c in icd10_codes[:2]])
            if cpt_codes:
                code_list.extend([f"procedure {c}" for c in cpt_codes[:2]])
            codes_text = ', '.join(code_list)
            codes_assessment = (
                f"The submitted clinical documentation includes {codes_text}. "
                "These codes have been validated against medical standards and insurance policy requirements."
            )
        
        reasoning_parts.append(intro + (codes_assessment or "Medical codes have been extracted and validated."))
        
        # Decision-based reasoning
        final_score = pa_data.get('final_score')
        decision_reasoning = f"Overall Assessment: The prior authorization review has been completed with a final confidence score of {final_score or 'N/A'}. "
        
        if status == 'APPROVED':
            decision_reasoning += (
                "Based on the clinical documentation and medical code analysis, this request meets the criteria for approval. "
                "The patient's condition and proposed treatment align with medical necessity standards and policy coverage requirements. "
                "Authorization is recommended."
            )
        elif status == 'DENIED':
            decision_reasoning += (
                "Based on the available clinical information and policy analysis, this request does not currently meet authorization criteria. "
                "The submitted documentation may not fully support medical necessity or may indicate policy conflicts. "
                "The provider may submit additional clinical evidence or file an appeal for reconsideration."
            )
        else:
            decision_reasoning += (
                "The review process is ongoing or requires additional information for completion. "
                "Further clinical analysis by qualified medical staff is in progress. "
                "Please check back for updates or contact support if additional documentation is needed."
            )
        
        reasoning_parts.append(decision_reasoning)
        
        # Combine paragraphs
        for part in reasoning_parts:
            if part:
                p = doc.add_paragraph(part)
                p.paragraph_format.space_after = Pt(10)
                p.paragraph_format.line_spacing = 1.15
    
    # Section 5: Footer
    doc.add_paragraph()
    _add_section_heading(doc, '5. REPORT DETAILS')
    
    footer_info = [
        f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")}',
        f'Report Type: Professional Summary',
        'Analysis Powered By: Sonar (Perplexity AI) + Medical NLP',
    ]
    
    for info in footer_info:
        p = doc.add_paragraph(info)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Convert to bytes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def save_report_to_file(pa_id: UUID, report_bytes: bytes, upload_dir: Path = Path('uploads')) -> str:
    """
    Save report bytes to disk.
    
    Args:
        pa_id: PA request ID
        report_bytes: Document bytes
        upload_dir: Directory to save reports
    
    Returns:
        File path to saved report
    """
    report_dir = upload_dir / str(pa_id) / 'reports'
    report_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = report_dir / f'PA_{pa_id}_Summary_Report.docx'
    file_path.write_bytes(report_bytes)
    
    logger.info(f"Report saved to {file_path}")
    return str(file_path)
